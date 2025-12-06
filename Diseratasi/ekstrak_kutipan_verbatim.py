from docx import Document
import re

doc = Document('Laporan_30_Artikel_Meditasi_Yogyakarta_2025.docx')

new_doc = Document()
new_doc.add_heading('Kutipan Verbatim dari 30 Artikel Meditasi Yogyakarta', 0)

for table in doc.tables:
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        
        no_match = re.search(r'^\d+$', cells[0]) if cells else None
        if no_match:
            no = cells[0]
            judul = cells[1] if len(cells) > 1 else ''
            
            verbatim_quotes = []
            for cell in cells:
                quotes = re.findall(r'"([^"]+)"', cell)
                verbatim_quotes.extend(quotes)
            
            if verbatim_quotes:
                new_doc.add_heading(f'Artikel {no}: {judul}', 2)
                for idx, quote in enumerate(verbatim_quotes, 1):
                    new_doc.add_paragraph(f'{idx}. "{quote}"', style='List Number')

new_doc.save('Laporan_30_Artikel_Meditasi_Yogyakarta_2025_REVISI2.docx')
print("File 'Laporan_30_Artikel_Meditasi_Yogyakarta_2025_REVISI2.docx' berhasil dibuat!")
