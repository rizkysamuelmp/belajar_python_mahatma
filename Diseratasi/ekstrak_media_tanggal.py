import json
from docx import Document
from collections import Counter

with open('data_artikel_meditasi_yogyakarta.json', 'r', encoding='utf-8') as f:
    data = json.load(f)

media_count = Counter([item['media'] for item in data])

doc = Document()
doc.add_heading('Daftar Media dan Tanggal Publikasi', 0)

doc.add_heading(f'Ringkasan Media (Total: {len(media_count)} media)', 1)
for media, count in sorted(media_count.items(), key=lambda x: x[1], reverse=True):
    doc.add_paragraph(f'{media}: {count} artikel')

doc.add_heading(f'Detail Media dan Tanggal (Total: {len(data)} artikel)', 1)
for item in data:
    doc.add_paragraph(f"{item['no']}. {item['media']} - {item['tanggal']}")

doc.save('Daftar_Media_Tanggal.docx')
print(f"File 'Daftar_Media_Tanggal.docx' berhasil dibuat!")
print(f"Total artikel: {len(data)}")
print(f"Total media: {len(media_count)}")
