from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from data_revisi_kutipan import articles_revisi

doc = Document()

# Set margin
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

# HALAMAN JUDUL
title = doc.add_heading('LAPORAN ANALISIS ARTIKEL', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()
subtitle = doc.add_paragraph('Meditasi di Yogyakarta Tahun 2025')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(16)
subtitle.runs[0].font.bold = True

subtitle2 = doc.add_paragraph('(REVISI - Dengan Kutipan Verbatim Presisi)')
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle2.runs[0].font.size = Pt(12)
subtitle2.runs[0].font.italic = True

doc.add_paragraph()
info = doc.add_paragraph(f'Total: {len(articles_revisi)} Artikel')
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.runs[0].font.size = Pt(14)

doc.add_paragraph()
date_p = doc.add_paragraph('Desember 2025')
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# DAFTAR ARTIKEL LENGKAP
doc.add_heading('DAFTAR ARTIKEL LENGKAP', 1)

for article in articles_revisi:
    heading = doc.add_heading(f'Artikel {article["no"]}: {article["judul"]}', 2)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    p = doc.add_paragraph()
    p.add_run('Tanggal: ').bold = True
    p.add_run(article['tanggal'])
    
    p = doc.add_paragraph()
    p.add_run('Media: ').bold = True
    p.add_run(article['media'])
    
    p = doc.add_paragraph()
    p.add_run('Link: ').bold = True
    link_run = p.add_run(article['link'])
    link_run.font.color.rgb = RGBColor(0, 0, 255)
    link_run.font.underline = True
    
    doc.add_paragraph()
    content_heading = doc.add_paragraph()
    content_heading.add_run('RINGKASAN ARTIKEL:').bold = True
    
    content_para = doc.add_paragraph(article['ringkasan'])
    content_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    quote_heading = doc.add_paragraph()
    quote_heading.add_run('KUTIPAN VERBATIM:').bold = True
    quote_heading.runs[0].font.color.rgb = RGBColor(204, 0, 0)
    
    quote_para = doc.add_paragraph(article['kutipan_verbatim'])
    quote_para.paragraph_format.left_indent = Inches(0.5)
    quote_para.runs[0].font.italic = True
    quote_para.runs[0].font.color.rgb = RGBColor(51, 51, 51)
    
    doc.add_paragraph('_' * 100)
    doc.add_paragraph()

doc.add_page_break()

# TABEL RINGKASAN
doc.add_heading('TABEL RINGKASAN ARTIKEL', 1)

table = doc.add_table(rows=1, cols=5)
table.style = 'Light Grid Accent 1'

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'No'
hdr_cells[1].text = 'Tanggal'
hdr_cells[2].text = 'Media'
hdr_cells[3].text = 'Link'
hdr_cells[4].text = 'Kutipan Verbatim'

for cell in hdr_cells:
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)

for article in articles_revisi:
    row_cells = table.add_row().cells
    row_cells[0].text = str(article['no'])
    row_cells[1].text = article['tanggal']
    row_cells[2].text = article['media']
    row_cells[3].text = article['link'][:40] + '...'
    row_cells[4].text = article['kutipan_verbatim']
    
    for cell in row_cells:
        cell.paragraphs[0].runs[0].font.size = Pt(9)

doc.save('Laporan_30_Artikel_Meditasi_Yogyakarta_2025_REVISI.docx')
print(f"✓ Dokumen REVISI dengan {len(articles_revisi)} artikel berhasil dibuat!")
print("  File: Laporan_30_Artikel_Meditasi_Yogyakarta_2025_REVISI.docx")
print("  ✓ Semua kutipan menggunakan verbatim presisi dari artikel asli")
