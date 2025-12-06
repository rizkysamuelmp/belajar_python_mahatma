import openpyxl
from docx import Document
import requests
from bs4 import BeautifulSoup
import time
from urllib.parse import urlparse

# Baca file Excel
wb = openpyxl.load_workbook('source_data/Olivia Rodrigo/Bank Link.xlsx')
ws = wb.active

# Buat dokumen Word
doc = Document()
doc.add_heading('Artikel dari Bank Link', 0)

# Iterasi setiap baris di Excel
counter = 1
for row in ws.iter_rows(min_row=2, values_only=True):
    if not row[0]:
        continue
    
    link = str(row[0]).strip()
    print(f"[{counter}] Mengambil: {link}")
    
    try:
        headers = {'User-Agent': 'Mozilla/5.0'}
        response = requests.get(link, timeout=10, headers=headers)
        response.raise_for_status()
        
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Ambil nama media dari domain
        domain = urlparse(link).netloc.replace('www.', '')
        
        # Cari tanggal publikasi
        date = None
        date_selectors = ['time', 'meta[property="article:published_time"]', 'meta[name="publish-date"]', '.date', '.published-date']
        for selector in date_selectors:
            elem = soup.select_one(selector)
            if elem:
                date = elem.get('datetime') or elem.get('content') or elem.get_text(strip=True)
                break
        
        # Hapus script dan style
        for tag in soup(['script', 'style', 'nav', 'footer', 'header']):
            tag.decompose()
        
        # Ambil judul
        title = soup.find('h1')
        title_text = title.get_text(strip=True) if title else "Tanpa Judul"
        
        # Ambil paragraf artikel
        paragraphs = soup.find_all('p')
        content = '\n\n'.join([p.get_text(strip=True) for p in paragraphs if len(p.get_text(strip=True)) > 50])
        
        # Tambahkan ke Word
        doc.add_heading(f"{counter}. {title_text}", 1)
        doc.add_paragraph(f"Media: {domain}")
        doc.add_paragraph(f"Tanggal: {date if date else 'Tidak ditemukan'}")
        doc.add_paragraph(f"Sumber: {link}")
        doc.add_paragraph("")
        doc.add_paragraph(content[:5000])
        doc.add_page_break()
        
        print(f"✓ Berhasil: {title_text}")
        counter += 1
        time.sleep(1)
        
    except Exception as e:
        print(f"✗ Skip: {link} - {str(e)}")
        continue

# Simpan dokumen
doc.save('Artikel_Bank_Link.docx')
print(f"\n✓ Selesai! Total {counter-1} artikel disimpan sebagai 'Artikel_Bank_Link.docx'")
