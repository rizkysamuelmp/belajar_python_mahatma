import pandas as pd
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from collections import Counter
from wordcloud import WordCloud
import matplotlib.pyplot as plt
from datetime import datetime
import os

TEMA_1 = "Revival Memory: Reproduksi gaya Y2K dalam Media Online"
TEMA_2 = "Identitas Generasional: Figur Selebriti Olivia Rodrigo sebagai Cultural Intermediary Y2K"

keywords_tema1 = ['y2k', 'nostalgia', '2000', 'retro', 'vintage', 'throwback', 'revival', 'aesthetic', 
                  'fashion', 'style', 'trend', 'era', 'millennium', 'noughties', 'early 2000s', '2000s']

keywords_tema2 = ['olivia rodrigo', 'generation', 'gen z', 'identity', 'cultural', 'influence', 
                  'celebrity', 'icon', 'represent', 'youth', 'generational', 'intermediary', 'young']

def scrape_artikel(url, timeout=10):
    try:
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
        response = requests.get(url, headers=headers, timeout=timeout)
        response.raise_for_status()
        soup = BeautifulSoup(response.content, 'html.parser')
        
        for script in soup(["script", "style"]):
            script.decompose()
        
        text = soup.get_text()
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = ' '.join(chunk for chunk in chunks if chunk)
        
        title = soup.find('title')
        title = title.get_text() if title else "No Title"
        
        return {'success': True, 'text': text, 'title': title}
    except Exception as e:
        return {'success': False, 'error': str(e)}

def cari_kutipan_relevan(text, keywords, context_length=150):
    text_lower = text.lower()
    kutipan = []
    
    for keyword in keywords:
        pattern = re.compile(r'.{0,' + str(context_length) + r'}' + re.escape(keyword.lower()) + 
                           r'.{0,' + str(context_length) + r'}', re.IGNORECASE)
        matches = pattern.findall(text)
        
        for match in matches[:2]:
            match_clean = ' '.join(match.strip().split())
            if len(match_clean) > 50 and match_clean not in kutipan:
                kutipan.append(match_clean)
    
    return kutipan[:3]

def ekstrak_nama_media(url):
    try:
        domain = url.split('/')[2].replace('www.', '')
        return domain.split('.')[0].upper()
    except:
        return "Unknown"

def ekstrak_tanggal_advanced(text, url, soup=None):
    """Ekstrak tanggal dengan berbagai metode"""
    # 1. Coba dari meta tags jika soup tersedia
    if soup:
        meta_tags = ['article:published_time', 'datePublished', 'publishdate', 'date']
        for tag in meta_tags:
            meta = soup.find('meta', property=tag) or soup.find('meta', {'name': tag})
            if meta and meta.get('content'):
                date_str = meta.get('content')
                try:
                    date_obj = datetime.fromisoformat(date_str.replace('Z', '+00:00'))
                    return date_obj.strftime('%Y-%m-%d')
                except:
                    pass
    
    # 2. Pattern tanggal dari teks
    patterns = [
        (r'(\d{4})-(\d{2})-(\d{2})', lambda m: f"{m.group(1)}-{m.group(2)}-{m.group(3)}"),
        (r'(\d{1,2})/(\d{1,2})/(\d{4})', lambda m: f"{m.group(3)}-{m.group(1).zfill(2)}-{m.group(2).zfill(2)}"),
        (r'(January|February|March|April|May|June|July|August|September|October|November|December)\s+(\d{1,2}),?\s+(\d{4})', 
         lambda m: convert_month_to_date(m.group(1), m.group(2), m.group(3))),
    ]
    
    for pattern, formatter in patterns:
        match = re.search(pattern, text[:3000])
        if match:
            try:
                return formatter(match)
            except:
                pass
    
    # 3. Dari URL
    url_patterns = [
        r'/(\d{4})/(\d{2})/(\d{2})/',
        r'/(\d{4})-(\d{2})-(\d{2})',
    ]
    
    for pattern in url_patterns:
        match = re.search(pattern, url)
        if match:
            return f"{match.group(1)}-{match.group(2)}-{match.group(3)}"
    
    return "Tidak tersedia"

def convert_month_to_date(month, day, year):
    months = {'January': '01', 'February': '02', 'March': '03', 'April': '04', 
              'May': '05', 'June': '06', 'July': '07', 'August': '08',
              'September': '09', 'October': '10', 'November': '11', 'December': '12'}
    return f"{year}-{months[month]}-{day.zfill(2)}"

def scrape_artikel_lengkap(url, timeout=10):
    """Scraping dengan ekstraksi tanggal yang lebih baik"""
    try:
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
        response = requests.get(url, headers=headers, timeout=timeout)
        response.raise_for_status()
        soup = BeautifulSoup(response.content, 'html.parser')
        
        for script in soup(["script", "style"]):
            script.decompose()
        
        text = soup.get_text()
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = ' '.join(chunk for chunk in chunks if chunk)
        
        title = soup.find('title')
        title = title.get_text() if title else "No Title"
        
        tanggal = ekstrak_tanggal_advanced(text, url, soup)
        
        return {'success': True, 'text': text, 'title': title, 'tanggal': tanggal}
    except Exception as e:
        return {'success': False, 'error': str(e)}

def buat_wordcloud(text, title, output_path):
    wordcloud = WordCloud(width=1600, height=800, background_color='white', 
                         colormap='viridis', max_words=80, relative_scaling=0.5).generate(text)
    
    plt.figure(figsize=(16, 8))
    plt.imshow(wordcloud, interpolation='bilinear')
    plt.axis('off')
    plt.title(title, fontsize=18, pad=20, fontweight='bold')
    plt.tight_layout(pad=0)
    plt.savefig(output_path, dpi=300, bbox_inches='tight')
    plt.close()

def buat_diagram_batang(all_kutipan, output_path):
    all_text = ' '.join(all_kutipan).lower()
    
    kategori = {
        'Fashion & Style': ['fashion', 'style', 'outfit', 'dress', 'clothes', 'aesthetic', 'wear'],
        'Music & Performance': ['music', 'song', 'album', 'concert', 'performance', 'tour', 'sing'],
        'Y2K & Nostalgia': ['y2k', 'nostalgia', '2000s', 'retro', 'throwback', 'vintage', 'early 2000'],
        'Generation Z': ['gen z', 'generation', 'youth', 'young', 'teenager', 'millennial'],
        'Cultural Impact': ['culture', 'influence', 'icon', 'trend', 'phenomenon', 'impact'],
        'Social Media': ['instagram', 'tiktok', 'social media', 'viral', 'online', 'internet'],
        'Celebrity & Fame': ['celebrity', 'star', 'famous', 'olivia rodrigo', 'artist', 'singer'],
        'Beauty & Makeup': ['beauty', 'makeup', 'hair', 'look', 'glam', 'cosmetic']
    }
    
    counts = {kat: sum(all_text.count(kw) for kw in keywords) for kat, keywords in kategori.items()}
    sorted_counts = dict(sorted(counts.items(), key=lambda x: x[1], reverse=True)[:8])
    
    plt.figure(figsize=(14, 7))
    bars = plt.bar(range(len(sorted_counts)), list(sorted_counts.values()), 
                   color='steelblue', edgecolor='navy', linewidth=1.5)
    plt.xlabel('Kategori Pembahasan', fontsize=13, fontweight='bold')
    plt.ylabel('Frekuensi Kemunculan', fontsize=13, fontweight='bold')
    plt.title('Pembahasan yang Sering Muncul dalam Artikel', fontsize=16, pad=20, fontweight='bold')
    plt.xticks(range(len(sorted_counts)), list(sorted_counts.keys()), rotation=45, ha='right', fontsize=11)
    plt.grid(axis='y', alpha=0.3, linestyle='--')
    
    for i, bar in enumerate(bars):
        height = bar.get_height()
        plt.text(bar.get_x() + bar.get_width()/2., height,
                f'{int(height)}', ha='center', va='bottom', fontsize=11, fontweight='bold')
    
    plt.tight_layout()
    plt.savefig(output_path, dpi=300, bbox_inches='tight')
    plt.close()

def buat_dokumen_word(data_hasil, ringkasan_url, ringkasan_media, output_path):
    doc = Document()
    
    # Judul
    title = doc.add_heading('Analisis Artikel Olivia Rodrigo - Tema Y2K', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph(f'Tanggal Analisis: {datetime.now().strftime("%d %B %Y")}')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # Tema
    doc.add_heading('Tema Analisis:', level=1)
    doc.add_paragraph(f'1. {TEMA_1}')
    doc.add_paragraph(f'2. {TEMA_2}')
    doc.add_paragraph()
    
    # Ringkasan Statistik
    doc.add_heading('Ringkasan Statistik', level=1)
    doc.add_paragraph(f"Total URL diproses: {ringkasan_media['total_url']}")
    doc.add_paragraph(f"URL berhasil diakses: {ringkasan_media['sukses']}")
    doc.add_paragraph(f"URL gagal diakses: {ringkasan_media['gagal']}")
    doc.add_paragraph(f"Jumlah media berbeda: {ringkasan_media['jumlah_media']}")
    doc.add_paragraph(f"Daftar media: {', '.join(ringkasan_media['daftar_media'])}")
    doc.add_paragraph()
    
    # Tabel Ringkasan URL
    doc.add_heading('Tabel Ringkasan URL yang Diproses', level=1)
    
    table_url = doc.add_table(rows=1, cols=5)
    table_url.style = 'Light Grid Accent 1'
    
    header_cells = table_url.rows[0].cells
    headers = ['No', 'URL', 'Nama Media', 'Tanggal', 'Keterangan']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
    
    for item in ringkasan_url:
        row_cells = table_url.add_row().cells
        row_cells[0].text = str(item['no'])
        row_cells[1].text = item['url'][:60] + '...' if len(item['url']) > 60 else item['url']
        row_cells[2].text = item['media']
        row_cells[3].text = item['tanggal']
        row_cells[4].text = item['status']
        
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    
    doc.add_page_break()
    
    # Tabel Kutipan
    doc.add_heading('Tabel Kutipan Relevan per Artikel', level=1)
    
    table_kutipan = doc.add_table(rows=1, cols=5)
    table_kutipan.style = 'Light Grid Accent 1'
    
    header_cells = table_kutipan.rows[0].cells
    headers = ['No', 'Nama Media', 'Tanggal', 'Kutipan Tema 1\n(Revival Memory Y2K)', 'Kutipan Tema 2\n(Identitas Generasional)']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    
    for item in data_hasil:
        row_cells = table_kutipan.add_row().cells
        row_cells[0].text = str(item['no'])
        row_cells[1].text = item['media']
        row_cells[2].text = item['tanggal']
        
        # Format kutipan tema 1
        if item['kutipan_tema1']:
            row_cells[3].text = '\n\n'.join([f"• {k}" for k in item['kutipan_tema1']])
        else:
            row_cells[3].text = "Tidak ada kutipan relevan"
        
        # Format kutipan tema 2
        if item['kutipan_tema2']:
            row_cells[4].text = '\n\n'.join([f"• {k}" for k in item['kutipan_tema2']])
        else:
            row_cells[4].text = "Tidak ada kutipan relevan"
        
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                paragraph.space_after = Pt(6)
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    
    doc.add_page_break()
    
    # Visualisasi
    doc.add_heading('Visualisasi Data', level=1)
    
    # Word Cloud Tema 1
    if os.path.exists('wordcloud_tema1_revisi.png'):
        doc.add_heading('Word Cloud - Tema 1: Revival Memory Y2K', level=2)
        doc.add_picture('wordcloud_tema1_revisi.png', width=Inches(6.5))
        doc.add_paragraph()
    
    # Word Cloud Tema 2
    if os.path.exists('wordcloud_tema2_revisi.png'):
        doc.add_heading('Word Cloud - Tema 2: Identitas Generasional', level=2)
        doc.add_picture('wordcloud_tema2_revisi.png', width=Inches(6.5))
        doc.add_paragraph()
    
    # Diagram Batang
    if os.path.exists('diagram_pembahasan_revisi.png'):
        doc.add_heading('Diagram Batang - Pembahasan yang Sering Muncul', level=2)
        doc.add_picture('diagram_pembahasan_revisi.png', width=Inches(6.5))
    
    doc.save(output_path)

# Main execution
print("="*70)
print("MEMULAI ANALISIS ARTIKEL OLIVIA RODRIGO (REVISI)")
print("="*70)

excel_path = 'Bank Link.xlsx'
df = pd.read_excel(excel_path)

print(f"\nDitemukan {len(df)} URL dalam file Excel")

url_column = None
date_column = None

for col in df.columns:
    if 'url' in col.lower() or 'link' in col.lower():
        url_column = col
    if 'date' in col.lower() or 'tanggal' in col.lower():
        date_column = col

if url_column is None:
    url_column = df.columns[0]

print(f"Kolom URL: {url_column}")
print(f"Kolom Tanggal: {date_column if date_column else 'Tidak ada (akan diekstrak dari artikel)'}")

data_hasil = []
ringkasan_url = []
all_text_tema1 = []
all_text_tema2 = []
all_kutipan = []

for idx, row in df.iterrows():
    url = row[url_column]
    no = idx + 1
    
    print(f"\n[{no}/{len(df)}] {url[:70]}...")
    
    result = scrape_artikel_lengkap(url)
    
    if result['success']:
        text = result['text']
        nama_media = ekstrak_nama_media(url)
        
        # Gunakan tanggal dari Excel jika ada, jika tidak ekstrak dari artikel
        if date_column and pd.notna(row[date_column]):
            tanggal = str(row[date_column])
            if 'Timestamp' in str(type(row[date_column])):
                tanggal = row[date_column].strftime('%Y-%m-%d')
        else:
            tanggal = result.get('tanggal', 'Tidak tersedia')
        
        kutipan_tema1 = cari_kutipan_relevan(text, keywords_tema1)
        kutipan_tema2 = cari_kutipan_relevan(text, keywords_tema2)
        
        if kutipan_tema1 or kutipan_tema2:
            data_hasil.append({
                'no': no,
                'media': nama_media,
                'tanggal': tanggal,
                'kutipan_tema1': kutipan_tema1,
                'kutipan_tema2': kutipan_tema2
            })
            
            if kutipan_tema1:
                all_text_tema1.append(' '.join(kutipan_tema1))
            if kutipan_tema2:
                all_text_tema2.append(' '.join(kutipan_tema2))
            
            all_kutipan.extend(kutipan_tema1 + kutipan_tema2)
            
            print(f"  ✓ Sukses | Tema 1: {len(kutipan_tema1)} kutipan | Tema 2: {len(kutipan_tema2)} kutipan | Tanggal: {tanggal}")
        else:
            print(f"  ✓ Sukses diakses, tidak ada kutipan relevan")
        
        ringkasan_url.append({
            'no': no,
            'url': url,
            'media': nama_media,
            'tanggal': tanggal,
            'status': 'Sukses'
        })
    else:
        print(f"  ✗ Gagal: {result['error'][:60]}")
        ringkasan_url.append({
            'no': no,
            'url': url,
            'media': 'N/A',
            'tanggal': 'N/A',
            'status': f"Gagal"
        })

sukses_urls = [r for r in ringkasan_url if r['status'] == 'Sukses']
media_list = list(set([r['media'] for r in sukses_urls]))

ringkasan_media = {
    'total_url': len(df),
    'sukses': len(sukses_urls),
    'gagal': len(df) - len(sukses_urls),
    'jumlah_media': len(media_list),
    'daftar_media': sorted(media_list)
}

print("\n" + "="*70)
print("RINGKASAN AKHIR:")
print("="*70)
print(f"Total URL: {ringkasan_media['total_url']}")
print(f"Berhasil: {ringkasan_media['sukses']}")
print(f"Gagal: {ringkasan_media['gagal']}")
print(f"Jumlah media: {ringkasan_media['jumlah_media']}")
print(f"Artikel dengan kutipan relevan: {len(data_hasil)}")
print("="*70)

# Buat visualisasi
if all_text_tema1:
    print("\n[1/4] Membuat word cloud tema 1...")
    combined_text_tema1 = ' '.join(all_text_tema1)
    buat_wordcloud(combined_text_tema1, 'Word Cloud - Tema 1: Revival Memory Y2K', 'wordcloud_tema1_revisi.png')
    print("      ✓ Selesai")

if all_text_tema2:
    print("[2/4] Membuat word cloud tema 2...")
    combined_text_tema2 = ' '.join(all_text_tema2)
    buat_wordcloud(combined_text_tema2, 'Word Cloud - Tema 2: Identitas Generasional', 'wordcloud_tema2_revisi.png')
    print("      ✓ Selesai")

if all_kutipan:
    print("[3/4] Membuat diagram batang...")
    buat_diagram_batang(all_kutipan, 'diagram_pembahasan_revisi.png')
    print("      ✓ Selesai")

print("[4/4] Membuat dokumen Word...")
output_path = 'Analisis_Tema_Olivia_Rodrigo_Revisi.docx'
buat_dokumen_word(data_hasil, ringkasan_url, ringkasan_media, output_path)
print(f"      ✓ Selesai: {output_path}")

print("\n" + "="*70)
print("✓ PROSES SELESAI! Semua file revisi telah dibuat.")
print("="*70)
