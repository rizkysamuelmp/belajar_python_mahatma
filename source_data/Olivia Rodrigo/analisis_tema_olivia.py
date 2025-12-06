import pandas as pd
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from collections import Counter
from wordcloud import WordCloud
import matplotlib.pyplot as plt
import numpy as np
from datetime import datetime

# Tema yang dicari
TEMA_1 = "Revival Memory: Reproduksi gaya Y2K dalam Media Online"
TEMA_2 = "Identitas Generasional: Figur Selebriti Olivia Rodrigo sebagai Cultural Intermediary Y2K"

# Keywords untuk tema 1
keywords_tema1 = ['y2k', 'nostalgia', '2000', 'retro', 'vintage', 'throwback', 'revival', 'aesthetic', 
                  'fashion', 'style', 'trend', 'era', 'millennium', 'noughties', 'early 2000s']

# Keywords untuk tema 2
keywords_tema2 = ['olivia rodrigo', 'generation', 'gen z', 'identity', 'cultural', 'influence', 
                  'celebrity', 'icon', 'represent', 'youth', 'generational', 'intermediary']

def baca_excel(file_path):
    """Membaca file Excel dan mengambil URL"""
    df = pd.read_excel(file_path)
    return df

def scrape_artikel(url, timeout=10):
    """Scraping artikel dari URL"""
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
        response = requests.get(url, headers=headers, timeout=timeout)
        response.raise_for_status()
        
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Hapus script dan style
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Ambil teks
        text = soup.get_text()
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = ' '.join(chunk for chunk in chunks if chunk)
        
        # Coba ambil judul
        title = soup.find('title')
        title = title.get_text() if title else "No Title"
        
        return {'success': True, 'text': text, 'title': title}
    except Exception as e:
        return {'success': False, 'error': str(e)}

def cari_kutipan_relevan(text, keywords, context_length=200):
    """Mencari kutipan relevan berdasarkan keywords"""
    text_lower = text.lower()
    kutipan = []
    
    for keyword in keywords:
        pattern = re.compile(r'.{0,' + str(context_length) + r'}' + re.escape(keyword.lower()) + 
                           r'.{0,' + str(context_length) + r'}', re.IGNORECASE)
        matches = pattern.findall(text)
        
        for match in matches[:2]:  # Ambil max 2 kutipan per keyword
            match_clean = match.strip()
            if len(match_clean) > 50 and match_clean not in kutipan:
                kutipan.append(match_clean)
    
    return kutipan[:3]  # Max 3 kutipan per tema

def ekstrak_nama_media(url):
    """Ekstrak nama media dari URL"""
    try:
        domain = url.split('/')[2]
        domain = domain.replace('www.', '')
        return domain.split('.')[0].upper()
    except:
        return "Unknown"

def ekstrak_tanggal(text, url):
    """Mencoba ekstrak tanggal dari teks atau URL"""
    # Pattern tanggal umum
    patterns = [
        r'\d{1,2}[-/]\d{1,2}[-/]\d{4}',
        r'\d{4}[-/]\d{1,2}[-/]\d{1,2}',
        r'(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}'
    ]
    
    for pattern in patterns:
        match = re.search(pattern, text[:2000])
        if match:
            return match.group()
    
    # Coba dari URL
    url_date = re.search(r'(\d{4})/(\d{2})/(\d{2})', url)
    if url_date:
        return f"{url_date.group(1)}-{url_date.group(2)}-{url_date.group(3)}"
    
    return "N/A"

def buat_wordcloud(all_text, output_path):
    """Membuat word cloud dari semua teks"""
    wordcloud = WordCloud(width=1200, height=600, background_color='white', 
                         colormap='viridis', max_words=100).generate(all_text)
    
    plt.figure(figsize=(12, 6))
    plt.imshow(wordcloud, interpolation='bilinear')
    plt.axis('off')
    plt.title('Word Cloud - Kata yang Sering Muncul', fontsize=16, pad=20)
    plt.tight_layout()
    plt.savefig(output_path, dpi=300, bbox_inches='tight')
    plt.close()

def buat_diagram_batang(all_kutipan, output_path):
    """Membuat diagram batang pembahasan yang sering muncul"""
    # Gabungkan semua kutipan
    all_text = ' '.join(all_kutipan).lower()
    
    # Kategori pembahasan
    kategori = {
        'Fashion & Style': ['fashion', 'style', 'outfit', 'dress', 'clothes', 'aesthetic'],
        'Music & Performance': ['music', 'song', 'album', 'concert', 'performance', 'tour'],
        'Y2K & Nostalgia': ['y2k', 'nostalgia', '2000s', 'retro', 'throwback', 'vintage'],
        'Generation Z': ['gen z', 'generation', 'youth', 'young', 'teenager', 'millennial'],
        'Cultural Impact': ['culture', 'influence', 'icon', 'trend', 'phenomenon', 'impact'],
        'Social Media': ['instagram', 'tiktok', 'social media', 'viral', 'online', 'internet']
    }
    
    counts = {}
    for kat, keywords in kategori.items():
        count = sum(all_text.count(kw) for kw in keywords)
        counts[kat] = count
    
    # Sort dan ambil top 8
    sorted_counts = dict(sorted(counts.items(), key=lambda x: x[1], reverse=True)[:8])
    
    plt.figure(figsize=(12, 6))
    bars = plt.bar(range(len(sorted_counts)), list(sorted_counts.values()), color='steelblue')
    plt.xlabel('Kategori Pembahasan', fontsize=12)
    plt.ylabel('Frekuensi Kemunculan', fontsize=12)
    plt.title('Pembahasan yang Sering Muncul dalam Artikel', fontsize=14, pad=20)
    plt.xticks(range(len(sorted_counts)), list(sorted_counts.keys()), rotation=45, ha='right')
    
    # Tambahkan nilai di atas bar
    for i, bar in enumerate(bars):
        height = bar.get_height()
        plt.text(bar.get_x() + bar.get_width()/2., height,
                f'{int(height)}', ha='center', va='bottom', fontsize=10)
    
    plt.tight_layout()
    plt.savefig(output_path, dpi=300, bbox_inches='tight')
    plt.close()

def buat_dokumen_word(data_hasil, ringkasan_url, ringkasan_media, output_path):
    """Membuat dokumen Word dengan format yang diminta"""
    doc = Document()
    
    # Judul
    title = doc.add_heading('Analisis Artikel Olivia Rodrigo - Tema Y2K', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'Tanggal Analisis: {datetime.now().strftime("%d %B %Y")}')
    doc.add_paragraph()
    
    # Tema
    doc.add_heading('Tema Analisis:', level=1)
    doc.add_paragraph(f'1. {TEMA_1}')
    doc.add_paragraph(f'2. {TEMA_2}')
    doc.add_paragraph()
    
    # Tabel Ringkasan URL
    doc.add_heading('Ringkasan URL yang Diproses', level=1)
    
    table_url = doc.add_table(rows=1, cols=5)
    table_url.style = 'Light Grid Accent 1'
    
    header_cells = table_url.rows[0].cells
    headers = ['No', 'URL', 'Nama Media', 'Tanggal', 'Keterangan']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
    
    for item in ringkasan_url:
        row_cells = table_url.add_row().cells
        row_cells[0].text = str(item['no'])
        row_cells[1].text = item['url'][:50] + '...' if len(item['url']) > 50 else item['url']
        row_cells[2].text = item['media']
        row_cells[3].text = item['tanggal']
        row_cells[4].text = item['status']
    
    doc.add_paragraph()
    
    # Ringkasan Media
    doc.add_heading('Ringkasan Media', level=1)
    doc.add_paragraph(f"Total URL diproses: {ringkasan_media['total_url']}")
    doc.add_paragraph(f"URL berhasil: {ringkasan_media['sukses']}")
    doc.add_paragraph(f"URL gagal: {ringkasan_media['gagal']}")
    doc.add_paragraph(f"Jumlah media berbeda: {ringkasan_media['jumlah_media']}")
    doc.add_paragraph(f"Daftar media: {', '.join(ringkasan_media['daftar_media'])}")
    doc.add_paragraph()
    
    # Tabel Kutipan
    doc.add_heading('Kutipan Relevan per Artikel', level=1)
    
    table_kutipan = doc.add_table(rows=1, cols=5)
    table_kutipan.style = 'Light Grid Accent 1'
    
    header_cells = table_kutipan.rows[0].cells
    headers = ['No', 'Nama Media', 'Tanggal', 'Kutipan Tema 1', 'Kutipan Tema 2']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
    
    for item in data_hasil:
        row_cells = table_kutipan.add_row().cells
        row_cells[0].text = str(item['no'])
        row_cells[1].text = item['media']
        row_cells[2].text = item['tanggal']
        row_cells[3].text = '\n\n'.join([f"• {k}" for k in item['kutipan_tema1']])
        row_cells[4].text = '\n\n'.join([f"• {k}" for k in item['kutipan_tema2']])
    
    doc.add_page_break()
    
    # Visualisasi
    doc.add_heading('Visualisasi Data', level=1)
    
    # Word Cloud
    if os.path.exists('wordcloud_olivia.png'):
        doc.add_heading('Word Cloud - Kata yang Sering Muncul', level=2)
        doc.add_picture('wordcloud_olivia.png', width=Inches(6))
        doc.add_paragraph()
    
    # Diagram Batang
    if os.path.exists('diagram_pembahasan.png'):
        doc.add_heading('Diagram Batang - Pembahasan yang Sering Muncul', level=2)
        doc.add_picture('diagram_pembahasan.png', width=Inches(6))
    
    doc.save(output_path)

# Main execution
import os

print("Memulai analisis artikel Olivia Rodrigo...")

# Baca file Excel
excel_path = 'Bank Link.xlsx'
df = baca_excel(excel_path)

print(f"Ditemukan {len(df)} URL dalam file Excel")
print(f"Kolom dalam Excel: {df.columns.tolist()}")

# Identifikasi kolom URL
url_column = None
for col in df.columns:
    if 'url' in col.lower() or 'link' in col.lower():
        url_column = col
        break

if url_column is None:
    url_column = df.columns[0]  # Gunakan kolom pertama jika tidak ada kolom URL

print(f"Menggunakan kolom: {url_column}")

# Proses setiap URL
data_hasil = []
ringkasan_url = []
all_text = []
all_kutipan = []

for idx, row in df.iterrows():
    url = row[url_column]
    no = idx + 1
    
    print(f"\n[{no}/{len(df)}] Memproses: {url}")
    
    # Scrape artikel
    result = scrape_artikel(url)
    
    if result['success']:
        text = result['text']
        nama_media = ekstrak_nama_media(url)
        tanggal = ekstrak_tanggal(text, url)
        
        # Cari kutipan untuk tema 1
        kutipan_tema1 = cari_kutipan_relevan(text, keywords_tema1)
        
        # Cari kutipan untuk tema 2
        kutipan_tema2 = cari_kutipan_relevan(text, keywords_tema2)
        
        if kutipan_tema1 or kutipan_tema2:
            data_hasil.append({
                'no': no,
                'media': nama_media,
                'tanggal': tanggal,
                'kutipan_tema1': kutipan_tema1,
                'kutipan_tema2': kutipan_tema2
            })
            
            all_text.append(text)
            all_kutipan.extend(kutipan_tema1 + kutipan_tema2)
            
            print(f"✓ Berhasil - Ditemukan {len(kutipan_tema1)} kutipan tema 1, {len(kutipan_tema2)} kutipan tema 2")
        else:
            print(f"✓ Berhasil diakses, tapi tidak ada kutipan relevan")
        
        ringkasan_url.append({
            'no': no,
            'url': url,
            'media': nama_media,
            'tanggal': tanggal,
            'status': 'Sukses'
        })
    else:
        print(f"✗ Gagal: {result['error']}")
        ringkasan_url.append({
            'no': no,
            'url': url,
            'media': 'N/A',
            'tanggal': 'N/A',
            'status': f"Gagal: {result['error'][:50]}"
        })

# Hitung ringkasan media
sukses_urls = [r for r in ringkasan_url if r['status'] == 'Sukses']
media_list = list(set([r['media'] for r in sukses_urls]))

ringkasan_media = {
    'total_url': len(df),
    'sukses': len(sukses_urls),
    'gagal': len(df) - len(sukses_urls),
    'jumlah_media': len(media_list),
    'daftar_media': sorted(media_list)
}

print("\n" + "="*60)
print("RINGKASAN:")
print(f"Total URL: {ringkasan_media['total_url']}")
print(f"Berhasil: {ringkasan_media['sukses']}")
print(f"Gagal: {ringkasan_media['gagal']}")
print(f"Jumlah media: {ringkasan_media['jumlah_media']}")
print(f"Media: {', '.join(ringkasan_media['daftar_media'])}")
print("="*60)

# Buat visualisasi
if all_text:
    print("\nMembuat word cloud...")
    combined_text = ' '.join(all_text)
    buat_wordcloud(combined_text, 'wordcloud_olivia.png')
    print("✓ Word cloud selesai")

if all_kutipan:
    print("Membuat diagram batang...")
    buat_diagram_batang(all_kutipan, 'diagram_pembahasan.png')
    print("✓ Diagram batang selesai")

# Buat dokumen Word
print("\nMembuat dokumen Word...")
output_path = 'Analisis_Tema_Olivia_Rodrigo.docx'
buat_dokumen_word(data_hasil, ringkasan_url, ringkasan_media, output_path)
print(f"✓ Dokumen Word selesai: {output_path}")

print("\n✓ SELESAI! Semua file telah dibuat.")
