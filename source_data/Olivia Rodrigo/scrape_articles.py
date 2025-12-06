import pandas as pd
import requests
from bs4 import BeautifulSoup
from docx import Document
from datetime import datetime
import time
import re

def extract_article_content(url):
    try:
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        soup = BeautifulSoup(response.content, 'html.parser')
        for script in soup(["script", "style"]):
            script.decompose()
        text = soup.get_text()
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        return ' '.join(chunk for chunk in chunks if chunk)
    except:
        return None

def find_quotes(text, keywords):
    if not text:
        return []
    quotes = []
    text_lower = text.lower()
    for keyword in keywords:
        pos = text_lower.find(keyword.lower())
        if pos != -1:
            start = max(0, pos - 200)
            end = min(len(text), pos + len(keyword) + 200)
            quote = text[start:end].strip()
            if quote:
                quotes.append(quote)
                break
    return quotes

excel_file = '/home/mahatma/belajar_python_mahatma/source_data/Olivia Rodrigo/Bank Link.xlsx'
df = pd.read_excel(excel_file)

theme1_kw = ['Y2K', 'y2k', '2000an', 'nostalgia', 'retro', 'revival', 'early 2000s']
theme2_kw = ['Olivia Rodrigo', 'olivia rodrigo', 'Gen Z', 'generation']

results = []
url_summary = []
total = len(df)

print(f"Memproses {total} URL...\n")

for idx, row in df.iterrows():
    url = row.get('Bank Link')
    media = row.get('Media Name', 'Unknown')
    date = row.get('Date', 'Unknown')
    
    if pd.isna(url):
        continue
    
    print(f"[{idx+1}/{total}] {media}...", end=' ')
    
    content = extract_article_content(url)
    
    if content:
        q1 = find_quotes(content, theme1_kw)
        q2 = find_quotes(content, theme2_kw)
        
        results.append({
            'no': idx + 1,
            'media': media,
            'tanggal': date,
            'theme1': '\n'.join(q1) if q1 else 'Tidak ada kutipan relevan',
            'theme2': '\n'.join(q2) if q2 else 'Tidak ada kutipan relevan'
        })
        
        url_summary.append({'no': idx + 1, 'url': url, 'media': media, 'tanggal': date, 'status': 'Sukses'})
        print("✓")
    else:
        url_summary.append({'no': idx + 1, 'url': url, 'media': media, 'tanggal': date, 'status': 'Gagal'})
        print("✗")
    
    time.sleep(0.5)

print("\nMembuat dokumen Word...")
doc = Document()

doc.add_heading('Hasil Scraping Artikel - Analisis Tema Y2K dan Olivia Rodrigo', 0)

doc.add_heading('Ringkasan URL yang Diakses', 1)
t1 = doc.add_table(rows=1, cols=5)
t1.style = 'Light Grid Accent 1'
h = t1.rows[0].cells
h[0].text = 'No'
h[1].text = 'URL'
h[2].text = 'Nama Media'
h[3].text = 'Tanggal'
h[4].text = 'Keterangan'

for item in url_summary:
    r = t1.add_row().cells
    r[0].text = str(item['no'])
    r[1].text = str(item['url'])
    r[2].text = str(item['media'])
    r[3].text = str(item['tanggal'])
    r[4].text = item['status']

sukses = sum(1 for i in url_summary if i['status'] == 'Sukses')
gagal = len(url_summary) - sukses

doc.add_paragraph()
doc.add_heading('Ringkasan Statistik', 1)
doc.add_paragraph(f'Total URL: {len(url_summary)}')
doc.add_paragraph(f'Sukses: {sukses}')
doc.add_paragraph(f'Gagal: {gagal}')

if sukses > 0:
    doc.add_paragraph()
    doc.add_heading('Media yang Berhasil Diakses', 1)
    media_list = [i['media'] for i in url_summary if i['status'] == 'Sukses']
    media_count = {}
    for m in media_list:
        media_count[m] = media_count.get(m, 0) + 1
    for m, c in media_count.items():
        doc.add_paragraph(f'• {m}: {c} artikel', style='List Bullet')

doc.add_page_break()
doc.add_heading('Kutipan Relevan per Artikel', 1)

t2 = doc.add_table(rows=1, cols=5)
t2.style = 'Light Grid Accent 1'
h2 = t2.rows[0].cells
h2[0].text = 'No'
h2[1].text = 'Nama Media'
h2[2].text = 'Tanggal'
h2[3].text = 'Tema 1: Revival Memory Y2K'
h2[4].text = 'Tema 2: Identitas Generasional Olivia Rodrigo'

for res in results:
    r = t2.add_row().cells
    r[0].text = str(res['no'])
    r[1].text = str(res['media'])
    r[2].text = str(res['tanggal'])
    r[3].text = res['theme1']
    r[4].text = res['theme2']

output = '/home/mahatma/belajar_python_mahatma/source_data/Hasil_Scraping_Artikel.docx'
doc.save(output)

print(f"\n✓ Selesai! File: {output}")
print(f"Total artikel berhasil: {len(results)}")
