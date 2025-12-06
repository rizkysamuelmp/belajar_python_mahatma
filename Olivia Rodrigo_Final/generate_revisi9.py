import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from wordcloud import WordCloud
import matplotlib.pyplot as plt

# Read Excel
df = pd.read_excel('Kutipan Relavan1.xlsx')
df.columns = ['Media Name', 'Date', 'Nostalgic', 'Vintage', 'Retro', 'Revival', 'Gen Z']

# Create document
doc = Document()
title = doc.add_heading('Analisis Tipologi Kutipan', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Theme 1: Revival Memory
doc.add_heading('Tema 1: Revival Memory - Representasi Trend Y2K dalam Media Online', 1)

theme1_data = []
for idx, row in df.iterrows():
    for col in ['Nostalgic', 'Vintage', 'Retro', 'Revival']:
        if pd.notna(row[col]):
            theme1_data.append({
                'Media': str(row['Media Name']) if pd.notna(row['Media Name']) else '',
                'Date': str(row['Date']) if pd.notna(row['Date']) else '',
                'Quote': str(row[col])
            })

# Table 1
doc.add_paragraph('Tabel 1: Kutipan Verbatim - Revival Memory')
table1 = doc.add_table(rows=1, cols=4)
table1.style = 'Light Grid Accent 1'

hdr = table1.rows[0].cells
hdr[0].text = 'No'
hdr[1].text = 'Nama Media'
hdr[2].text = 'Tanggal'
hdr[3].text = 'Kutipan Verbatim'

for i, item in enumerate(theme1_data, 1):
    row = table1.add_row().cells
    row[0].text = str(i)
    row[1].text = item['Media']
    row[2].text = item['Date']
    row[3].text = item['Quote']

# Theme 2: Gen Z
doc.add_page_break()
doc.add_heading('Tema 2: Identitas Gen Z - Figur Selebriti Olivia Rodrigo Dalam Trend Y2K', 1)

theme2_data = []
for idx, row in df.iterrows():
    if pd.notna(row['Gen Z']):
        theme2_data.append({
            'Media': str(row['Media Name']) if pd.notna(row['Media Name']) else '',
            'Date': str(row['Date']) if pd.notna(row['Date']) else '',
            'Quote': str(row['Gen Z'])
        })

# Table 2
doc.add_paragraph('Tabel 2: Kutipan Verbatim - Identitas Gen Z')
table2 = doc.add_table(rows=1, cols=4)
table2.style = 'Light Grid Accent 1'

hdr = table2.rows[0].cells
hdr[0].text = 'No'
hdr[1].text = 'Nama Media'
hdr[2].text = 'Tanggal'
hdr[3].text = 'Kutipan Verbatim'

for i, item in enumerate(theme2_data, 1):
    row = table2.add_row().cells
    row[0].text = str(i)
    row[1].text = item['Media']
    row[2].text = item['Date']
    row[3].text = item['Quote']

# Word Clouds
doc.add_page_break()
doc.add_heading('Visualisasi Word Cloud', 1)

text1 = ' '.join([item['Quote'] for item in theme1_data])
wc1 = WordCloud(width=800, height=400, background_color='white').generate(text1)
plt.figure(figsize=(10, 5))
plt.imshow(wc1, interpolation='bilinear')
plt.axis('off')
plt.title('Revival Memory: Representasi Trend Y2K dalam Media Online')
plt.tight_layout()
plt.savefig('wordcloud_revival_memory.png', dpi=300, bbox_inches='tight')
plt.close()

doc.add_paragraph('Word Cloud 1: Revival Memory')
doc.add_picture('wordcloud_revival_memory.png', width=Inches(6))

text2 = ' '.join([item['Quote'] for item in theme2_data])
wc2 = WordCloud(width=800, height=400, background_color='white').generate(text2)
plt.figure(figsize=(10, 5))
plt.imshow(wc2, interpolation='bilinear')
plt.axis('off')
plt.title('Identitas Gen Z: Figur Selebriti Olivia Rodrigo Dalam Trend Y2K')
plt.tight_layout()
plt.savefig('wordcloud_gen_z.png', dpi=300, bbox_inches='tight')
plt.close()

doc.add_paragraph('Word Cloud 2: Identitas Gen Z')
doc.add_picture('wordcloud_gen_z.png', width=Inches(6))

# Summary
doc.add_page_break()
doc.add_heading('Kesimpulan', 1)

theme1_media = set([item['Media'] for item in theme1_data if item['Media']])
theme2_media = set([item['Media'] for item in theme2_data if item['Media']])
all_media = theme1_media.union(theme2_media)

summary = f"""
Ringkasan Analisis:

1. Tema Revival Memory (Nostalgic, Vintage, Retro, Revival):
   - Total kutipan: {len(theme1_data)}
   - Jumlah media: {len(theme1_media)}

2. Tema Identitas Gen Z:
   - Total kutipan: {len(theme2_data)}
   - Jumlah media: {len(theme2_media)}

3. Total keseluruhan:
   - Total media unik yang dianalisis: {len(all_media)}
   - Total kutipan: {len(theme1_data) + len(theme2_data)}
"""

doc.add_paragraph(summary)

doc.save('revisi9.docx')
print(f"Dokumen berhasil dibuat!")
print(f"Tema 1: {len(theme1_data)} kutipan dari {len(theme1_media)} media")
print(f"Tema 2: {len(theme2_data)} kutipan dari {len(theme2_media)} media")
print(f"Total media unik: {len(all_media)}")
