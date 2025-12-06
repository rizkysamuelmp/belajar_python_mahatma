import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from wordcloud import WordCloud
import matplotlib.pyplot as plt
from collections import Counter
import re

# Read Excel file
df = pd.read_excel('Kutipan Relavan1.xlsx')

# Clean column names
df.columns = ['Media Name', 'Date', 'Nostalgic', 'Vintage', 'Retro', 'Revival', 'Gen Z']

# Create document
doc = Document()

# Add title
title = doc.add_heading('Analisis Tipologi Kutipan', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Theme 1: Revival Memory
doc.add_heading('Tema 1: Revival Memory - Representasi Trend Y2K dalam Media Online', 1)

# Collect data for Theme 1 (Nostalgic, Vintage, Retro, Revival)
theme1_data = []
for idx, row in df.iterrows():
    media = row['Media Name']
    date = row['Date']
    
    for col in ['Nostalgic', 'Vintage', 'Retro', 'Revival']:
        if pd.notna(row[col]) and pd.notna(media):
            theme1_data.append({
                'Media': str(media),
                'Date': str(date) if pd.notna(date) else '',
                'Quote': str(row[col])
            })

# Create Table 1
doc.add_paragraph('Tabel 1: Kutipan Verbatim - Revival Memory')
table1 = doc.add_table(rows=1, cols=3)
table1.style = 'Light Grid Accent 1'

# Header row
hdr_cells = table1.rows[0].cells
hdr_cells[0].text = 'No'
hdr_cells[1].text = 'Nama Media'
hdr_cells[2].text = 'Tanggal'

# Add data
for i, item in enumerate(theme1_data, 1):
    row_cells = table1.add_row().cells
    row_cells[0].text = str(i)
    row_cells[1].text = item['Media']
    row_cells[2].text = item['Date']

# Theme 2: Identitas Gen Z
doc.add_page_break()
doc.add_heading('Tema 2: Identitas Gen Z - Figur Selebriti Olivia Rodrigo Dalam Trend Y2K', 1)

# Collect data for Theme 2 (Gen Z)
theme2_data = []
for idx, row in df.iterrows():
    media = row['Media Name']
    date = row['Date']
    
    if pd.notna(row['Gen Z']) and pd.notna(media):
        theme2_data.append({
            'Media': str(media),
            'Date': str(date) if pd.notna(date) else '',
            'Quote': str(row['Gen Z'])
        })

# Create Table 2
doc.add_paragraph('Tabel 2: Kutipan Verbatim - Identitas Gen Z')
table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Light Grid Accent 1'

# Header row
hdr_cells = table2.rows[0].cells
hdr_cells[0].text = 'No'
hdr_cells[1].text = 'Nama Media'
hdr_cells[2].text = 'Tanggal'

# Add data
for i, item in enumerate(theme2_data, 1):
    row_cells = table2.add_row().cells
    row_cells[0].text = str(i)
    row_cells[1].text = item['Media']
    row_cells[2].text = item['Date']

# Generate Word Clouds
doc.add_page_break()
doc.add_heading('Visualisasi Word Cloud', 1)

# Word Cloud 1: Revival Memory
text1 = ' '.join([item['Quote'] for item in theme1_data])
wordcloud1 = WordCloud(width=800, height=400, background_color='white').generate(text1)
plt.figure(figsize=(10, 5))
plt.imshow(wordcloud1, interpolation='bilinear')
plt.axis('off')
plt.title('Word Cloud - Revival Memory: Representasi Trend Y2K dalam Media Online')
plt.tight_layout()
plt.savefig('wordcloud_revival_memory.png', dpi=300, bbox_inches='tight')
plt.close()

doc.add_paragraph('Word Cloud 1: Revival Memory - Representasi Trend Y2K dalam Media Online')
doc.add_picture('wordcloud_revival_memory.png', width=Inches(6))

# Word Cloud 2: Identitas Gen Z
text2 = ' '.join([item['Quote'] for item in theme2_data])
wordcloud2 = WordCloud(width=800, height=400, background_color='white').generate(text2)
plt.figure(figsize=(10, 5))
plt.imshow(wordcloud2, interpolation='bilinear')
plt.axis('off')
plt.title('Word Cloud - Identitas Gen Z: Figur Selebriti Olivia Rodrigo Dalam Trend Y2K')
plt.tight_layout()
plt.savefig('wordcloud_gen_z.png', dpi=300, bbox_inches='tight')
plt.close()

doc.add_paragraph('Word Cloud 2: Identitas Gen Z - Figur Selebriti Olivia Rodrigo Dalam Trend Y2K')
doc.add_picture('wordcloud_gen_z.png', width=Inches(6))

# Summary
doc.add_page_break()
doc.add_heading('Kesimpulan', 1)

# Count unique media
theme1_media = set([item['Media'] for item in theme1_data])
theme2_media = set([item['Media'] for item in theme2_data])
all_media = theme1_media.union(theme2_media)

summary_text = f"""
Ringkasan Analisis:

1. Tema Revival Memory (Nostalgic, Vintage, Retro, Revival):
   - Total kutipan: {len(theme1_data)}
   - Jumlah media: {len(theme1_media)}

2. Tema Identitas Gen Z:
   - Total kutipan: {len(theme2_data)}
   - Jumlah media: {len(theme2_media)}

3. Total keseluruhan:
   - Total media unik yang dianalisis: {len(all_media)}
   - Total kutipan: {len(theme1_data) + len(theme2_data)}
"""

doc.add_paragraph(summary_text)

# Save document
doc.save('revisi4.docx')
print(f"Document created successfully!")
print(f"Theme 1 (Revival Memory): {len(theme1_data)} quotes from {len(theme1_media)} media")
print(f"Theme 2 (Gen Z Identity): {len(theme2_data)} quotes from {len(theme2_media)} media")
print(f"Total unique media: {len(all_media)}")
