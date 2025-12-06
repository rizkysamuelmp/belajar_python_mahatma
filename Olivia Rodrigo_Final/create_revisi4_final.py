import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from wordcloud import WordCloud
import matplotlib.pyplot as plt
from collections import Counter
import re

# Baca file Excel
df = pd.read_excel('Kutipan Relavan.xlsx')

print(f"Total baris di Excel: {len(df)}")

# Fungsi untuk ekstrak frasa 2-4 kata yang presisi
def extract_precise_phrases(text):
    """Ekstrak frasa 2-4 kata yang presisi dan relevan"""
    if pd.isna(text) or str(text).strip() == '':
        return []
    
    text = str(text).strip()
    words = text.split()
    phrases = []
    
    # Ekstrak frasa 2-4 kata
    for length in [4, 3, 2]:  # Prioritas 4 kata, lalu 3, lalu 2
        for i in range(len(words) - length + 1):
            phrase = ' '.join(words[i:i+length])
            # Filter frasa yang mengandung kata kunci penting
            phrase_lower = phrase.lower()
            if any(kw in phrase_lower for kw in ['nostalgic', 'vintage', 'retro', 'revival', 
                                                   'y2k', '2000s', 'trend', 'fashion', 
                                                   'gen z', 'generation', 'olivia rodrigo']):
                phrases.append(phrase)
    
    # Ambil frasa terbaik (yang paling pendek dan relevan)
    if phrases:
        return [min(phrases, key=len)]
    
    # Jika tidak ada frasa dengan kata kunci, ambil 3 kata pertama
    if len(words) >= 3:
        return [' '.join(words[:3])]
    elif len(words) >= 2:
        return [' '.join(words[:2])]
    
    return []

# Proses data untuk Tema 1: Revival Memory
revival_data = []
revival_text_all = []

revival_cols = {
    'Precise Verbatim "Nostalgic': 'Nostalgic',
    'Precise Verbatim "Vintage"': 'Vintage',
    'Precise Verbatim "Retro"': 'Retro',
    'Precise Verbatim "Revival"': 'Revival'
}

for idx, row in df.iterrows():
    media = row.get('Media Name', '')
    date = row.get('Date', '')
    
    if pd.notna(media) and str(media).strip() != '':
        for col, category in revival_cols.items():
            if col in df.columns and pd.notna(row[col]) and str(row[col]).strip() != '':
                text = str(row[col]).strip()
                phrases = extract_precise_phrases(text)
                
                if phrases:
                    revival_data.append({
                        'Media': str(media).strip(),
                        'Date': str(date).strip() if pd.notna(date) else '',
                        'Verbatim': phrases[0],
                        'Category': category,
                        'FullText': text
                    })
                    revival_text_all.append(text)

# Proses data untuk Tema 2: Identitas Gen Z
genz_data = []
genz_text_all = []

genz_col = 'Precise Verbatim "Gen Z"'

for idx, row in df.iterrows():
    media = row.get('Media Name', '')
    date = row.get('Date', '')
    
    if pd.notna(media) and str(media).strip() != '':
        if genz_col in df.columns and pd.notna(row[genz_col]) and str(row[genz_col]).strip() != '':
            text = str(row[genz_col]).strip()
            phrases = extract_precise_phrases(text)
            
            if phrases:
                genz_data.append({
                    'Media': str(media).strip(),
                    'Date': str(date).strip() if pd.notna(date) else '',
                    'Verbatim': phrases[0],
                    'FullText': text
                })
                genz_text_all.append(text)

# Hitung media unik
media_revival_list = [item['Media'] for item in revival_data]
media_genz_list = [item['Media'] for item in genz_data]

unique_media_revival = set(media_revival_list)
unique_media_genz = set(media_genz_list)
all_unique_media = unique_media_revival.union(unique_media_genz)

media_revival_count = Counter(media_revival_list)
media_genz_count = Counter(media_genz_list)
category_count = Counter([item['Category'] for item in revival_data])

print(f"\nKutipan Revival Memory: {len(revival_data)}")
print(f"Kutipan Gen Z: {len(genz_data)}")
print(f"Total Media Unik: {len(all_unique_media)}")

# Buat dokumen Word
doc = Document()

# Judul
title = doc.add_heading('Analisis Tipologi Kutipan Olivia Rodrigo', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# KESIMPULAN
doc.add_heading('Kesimpulan', 1)

p1 = doc.add_paragraph()
run1 = p1.add_run('Total Media yang Dianalisis: ')
run1.bold = True
run1.font.size = Pt(12)
run2 = p1.add_run(f'{len(all_unique_media)} media')
run2.font.size = Pt(12)

doc.add_paragraph()

# Revival Memory
p2 = doc.add_paragraph()
run3 = p2.add_run('Tema 1: Revival Memory - Representasi Trend Y2K dalam Media Online')
run3.bold = True
run3.font.size = Pt(12)

p3 = doc.add_paragraph()
run4 = p3.add_run(f'Total Kutipan: ')
run4.bold = True
run5 = p3.add_run(f'{len(revival_data)} kutipan (frasa 2-4 kata presisi) dari {len(unique_media_revival)} media')

p_cat = doc.add_paragraph('Rincian per kategori:', style='List Bullet')
for cat, count in sorted(category_count.items()):
    p = doc.add_paragraph(f'{cat}: {count} kutipan', style='List Bullet 2')

if media_revival_count:
    p_top = doc.add_paragraph('Media dengan kutipan terbanyak:', style='List Bullet')
    for media, count in media_revival_count.most_common(5):
        p = doc.add_paragraph(f'{media}: {count} kutipan', style='List Bullet 2')

doc.add_paragraph()

# Gen Z
p4 = doc.add_paragraph()
run6 = p4.add_run('Tema 2: Identitas Gen Z - Figur Selebriti Olivia Rodrigo Dalam Trend Y2K')
run6.bold = True
run6.font.size = Pt(12)

p5 = doc.add_paragraph()
run7 = p5.add_run(f'Total Kutipan: ')
run7.bold = True
run8 = p5.add_run(f'{len(genz_data)} kutipan (frasa 2-4 kata presisi) dari {len(unique_media_genz)} media')

if media_genz_count:
    p_top2 = doc.add_paragraph('Media dengan kutipan terbanyak:', style='List Bullet')
    for media, count in media_genz_count.most_common(5):
        p = doc.add_paragraph(f'{media}: {count} kutipan', style='List Bullet 2')

doc.add_paragraph()

doc.add_heading('Daftar Lengkap Media yang Dianalisis:', 2)
for i, media in enumerate(sorted(all_unique_media), 1):
    p = doc.add_paragraph(f'{i}. {media}', style='List Number')

doc.add_page_break()

# TABEL 1: Revival Memory
doc.add_heading('Tabel 1: Revival Memory - Representasi Trend Y2K dalam Media Online', 1)
doc.add_paragraph(f'Total: {len(revival_data)} kutipan (frasa 2-4 kata presisi)')

table1 = doc.add_table(rows=1, cols=4)
table1.style = 'Light Grid Accent 1'

hdr = table1.rows[0].cells
hdr[0].text = 'No'
hdr[1].text = 'Nama Media'
hdr[2].text = 'Tanggal'
hdr[3].text = 'Kutipan Verbatim (2-4 Kata Presisi)'

for cell in hdr:
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(10)

for i, item in enumerate(revival_data, 1):
    row = table1.add_row().cells
    row[0].text = str(i)
    row[1].text = item['Media']
    row[2].text = item['Date']
    row[3].text = item['Verbatim']
    
    for cell in row:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

doc.add_page_break()

# TABEL 2: Identitas Gen Z
doc.add_heading('Tabel 2: Identitas Gen Z - Figur Selebriti Olivia Rodrigo Dalam Trend Y2K', 1)
doc.add_paragraph(f'Total: {len(genz_data)} kutipan (frasa 2-4 kata presisi)')

table2 = doc.add_table(rows=1, cols=4)
table2.style = 'Light Grid Accent 1'

hdr2 = table2.rows[0].cells
hdr2[0].text = 'No'
hdr2[1].text = 'Nama Media'
hdr2[2].text = 'Tanggal'
hdr2[3].text = 'Kutipan Verbatim (2-4 Kata Presisi)'

for cell in hdr2:
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(10)

for i, item in enumerate(genz_data, 1):
    row = table2.add_row().cells
    row[0].text = str(i)
    row[1].text = item['Media']
    row[2].text = item['Date']
    row[3].text = item['Verbatim']
    
    for cell in row:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

doc.save('Revisi4.docx')
print("\n✓ File Word 'Revisi4.docx' berhasil dibuat")

# VISUALISASI
if revival_text_all:
    plt.figure(figsize=(16, 8))
    wordcloud1 = WordCloud(width=1600, height=800, background_color='white', 
                           colormap='Blues', max_words=200, relative_scaling=0.5,
                           min_font_size=10, collocations=False).generate(' '.join(revival_text_all))
    plt.imshow(wordcloud1, interpolation='bilinear')
    plt.axis('off')
    plt.title('Word Cloud: Revival Memory - Representasi Trend Y2K dalam Media Online', 
              fontsize=22, fontweight='bold', pad=20)
    plt.tight_layout()
    plt.savefig('Revisi4_WordCloud_Revival.png', dpi=300, bbox_inches='tight', facecolor='white')
    print("✓ Word Cloud Revival Memory berhasil dibuat")
    plt.close()

if genz_text_all:
    plt.figure(figsize=(16, 8))
    wordcloud2 = WordCloud(width=1600, height=800, background_color='white', 
                           colormap='Purples', max_words=200, relative_scaling=0.5,
                           min_font_size=10, collocations=False).generate(' '.join(genz_text_all))
    plt.imshow(wordcloud2, interpolation='bilinear')
    plt.axis('off')
    plt.title('Word Cloud: Identitas Gen Z - Figur Selebriti Olivia Rodrigo Dalam Trend Y2K', 
              fontsize=22, fontweight='bold', pad=20)
    plt.tight_layout()
    plt.savefig('Revisi4_WordCloud_GenZ.png', dpi=300, bbox_inches='tight', facecolor='white')
    print("✓ Word Cloud Identitas Gen Z berhasil dibuat")
    plt.close()

print("\n" + "="*75)
print("RINGKASAN ANALISIS")
print("="*75)
print(f"Total Media: {len(all_unique_media)} media")
print(f"\nRevival Memory: {len(revival_data)} kutipan dari {len(unique_media_revival)} media")
for cat, count in sorted(category_count.items()):
    print(f"  - {cat}: {count} kutipan")
print(f"\nIdentitas Gen Z: {len(genz_data)} kutipan dari {len(unique_media_genz)} media")
print("="*75)
print("\n✅ KUTIPAN VERBATIM 2-4 KATA PRESISI")
print("✅ KUTIPAN BERULANG TETAP DIMASUKKAN")
